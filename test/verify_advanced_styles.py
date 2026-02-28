import sys
from pathlib import Path
import io
from docx import Document

# Add project root to path
root_path = Path(__file__).parent.parent
sys.path.append(str(root_path))

from src.file_formatting.formatting import generate_report

def verify_ieee_numbering():
    print("Verifying IEEE Numbering...", end="")
    structure = [
        {"type": "heading", "text": "Introduction"},
        {"type": "heading", "text": "Methods"}
    ]
    output = io.BytesIO()
    generate_report(structure, output, style_name="IEEE")
    output.seek(0)
    
    doc = Document(output)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    
    if paragraphs[0].startswith("1. Introduction") and paragraphs[1].startswith("2. Methods"):
        print(" Success!")
        return True
    else:
        print(f" Failed. Got: {paragraphs}")
        return False

def verify_apa_hanging_indent():
    print("Verifying APA Hanging Indent...", end="")
    structure = [
        {"type": "reference", "text": "Smith, J. (2020). The Art of Coding."}
    ]
    output = io.BytesIO()
    generate_report(structure, output, style_name="APA")
    output.seek(0)
    
    doc = Document(output)
    # The reference is the only paragraph
    p = doc.paragraphs[0]
    
    # Check indentation properties
    # Note: python-docx stores indent values in English Metric Units (EMU) or None
    # -0.5 inches is approx -457200 EMU
    
    first_line = p.paragraph_format.first_line_indent
    left_indent = p.paragraph_format.left_indent
    
    if first_line is not None and first_line < 0 and left_indent is not None and left_indent > 0:
        print(" Success!")
        return True
    else:
        print(f" Failed. First Line: {first_line}, Left: {left_indent}")
        return False

if __name__ == "__main__":
    passed = True
    passed &= verify_ieee_numbering()
    passed &= verify_apa_hanging_indent()
    
    if passed:
        print("\nAll advanced checks passed!")
        sys.exit(0)
    else:
        print("\nSome checks failed.")
        sys.exit(1)
