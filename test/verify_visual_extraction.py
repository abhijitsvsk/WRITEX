import sys
import os
from docx import Document

# Add src to path
sys.path.append(os.getcwd())
from src.analysis.style_analyzer import StyleAnalyzer

def test_visual_extraction():
    print("Testing Visual Style Extraction...")
    
    # 1. Create a dummy DOCX with specific styles
    test_file = "test_style_source.docx"
    doc = Document()
    
    # Add a paragraph with specific font and spacing
    p = doc.add_paragraph("This is a test paragraph.")
    p.paragraph_format.line_spacing = 2.0
    for run in p.runs:
        run.font.name = "Courier New"
        run.font.size = 14 * 12700 # Approx conversion if needed, but docx usually uses Pt object or just sets text
    
    # Simpler way for python-docx to ensure property is set for reading
    from docx.shared import Pt
    run = p.add_run(" More text.")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    
    # Add another paragraph to weigh the stats
    p2 = doc.add_paragraph("Secondary text.")
    p2.paragraph_format.line_spacing = 1.5
    for run in p2.runs:
        run.font.name = "Arial"
        run.font.size = Pt(12)

    doc.save(test_file)
    print(f"Created temporary file: {test_file}")

    # 2. Analyze it
    analyzer = StyleAnalyzer()
    style_profile = analyzer.analyze_visual_style(test_file)
    
    print(f"Extracted Profile: {style_profile}")
    
    # 3. Assertions
    # We expect Arial and 12pt to be most common
    try:
        assert style_profile["font_name"] == "Arial", f"Expected Arial, got {style_profile['font_name']}"
        # Size might be float, check close enough
        assert style_profile["font_size"] == 12.0, f"Expected 12.0, got {style_profile['font_size']}"
        print("SUCCESS: Visual extraction verified.")
    except AssertionError as e:
        print(f"FAILURE: {e}")
    except Exception as e:
        print(f"ERROR: {e}")
        
    # Cleanup
    if os.path.exists(test_file):
        os.remove(test_file)

if __name__ == "__main__":
    test_visual_extraction()
