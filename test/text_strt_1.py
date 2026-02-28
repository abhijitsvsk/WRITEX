from docx import Document

doc = Document(r"C:\Users\jithu\OneDrive\Documents\ENGLISH Report.docx")   # open existing
doc = Document()               # create new
for i in doc.paragraphs:
    print(i.text)