import os
import io
import json
from src.core.compiler import DocumentCompiler
from src.analysis.code_analyzer import CodeAnalyzer
from src.file_formatting.formatting import generate_report
from src.validation.validator import DocumentValidator
from docx import Document

stress_ast = [
    {"type": "title", "text": "Stress Test Report"},
    {"type": "title_page_body", "text": "Mock Submission Data\nJohn Doe"},
    {"type": "section_header", "text": "CERTIFICATE"},
    {"type": "paragraph", "text": "This is to certify..."},
    {"type": "signature_block", "text": "Signatures"},
    {"type": "section_header", "text": "ACKNOWLEDGEMENT"},
    {"type": "paragraph", "text": "I thank my guide..."},
    {"type": "section_header", "text": "ABSTRACT"},
    {"type": "paragraph", "text": "This is the abstract paragraph."},
    {"type": "toc"},
    {"type": "lof"},
    {"type": "chapter", "text": "INTRODUCTION"},
    {"type": "paragraph", "text": "Introduction text. " * 50},
    {"type": "subheading", "text": "Background"},
    {"type": "paragraph", "text": "Background text. " * 50},
    {"type": "paragraph", "text": "    "}, # Whitespace drop test
    {"type": "paragraph", "text": "[Extract Code: mock_func]"}, # Extract Code tag drop test
    {"type": "figure", "caption": "System Architecture Diagram"},
    {"type": "chapter", "text": "METHODOLOGY"},
    {"type": "subheading", "text": "Design"},
    {"type": "paragraph", "text": "Design text. " * 85},
    {"type": "code_block", "text": "def mock_func():\n    pass\n"},
    {"type": "subsubheading", "text": "Algorithm"},
    {"type": "paragraph", "text": "Algorithm details. " * 120}
]

validator = DocumentValidator()
is_clean, errors, healed_ast = validator._run_pass(stress_ast)

print("Validator run complete.")
print(f"Errors caught: {errors}")

buf = io.BytesIO()
generate_report(healed_ast, buf)

buf.seek(0)
doc = Document(buf)

print(f"Generated DOCX successfully. Total Paragraphs: {len(doc.paragraphs)}")

# Let's inspect the TOC text to verify if Front-Matter pagination logic worked.
toc_paragraphs = [p.text for p in doc.paragraphs if '\t' in p.text and 'Chapter' not in p.text and 'Figure' not in p.text]
print("TOC Entries (Front Matter Check):")
for t in toc_paragraphs:
    print(t)

print("TOC Entries (Chapters Check):")
chapter_entries = [p.text for p in doc.paragraphs if '\t' in p.text and 'Chapter' in p.text]
for c in chapter_entries:
    print(c)
