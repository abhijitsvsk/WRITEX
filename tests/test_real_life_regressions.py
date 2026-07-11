import io

from docx import Document
from docx.shared import Inches, Pt

from src.analysis.style_analyzer import StyleAnalyzer
from src.file_formatting.formatting import generate_report, StyleConfig
from src.models.constraints import ResolvedConstraints
from src.validation.validator import DocumentValidator


def test_document_validator_accepts_app_constraints_argument():
    validator = DocumentValidator(constraints=ResolvedConstraints(font_size=14))

    structure = [
        {"type": "section_header", "text": "Certificate"},
        {"type": "paragraph", "text": "Certificate body."},
        {"type": "section_header", "text": "Acknowledgement"},
        {"type": "paragraph", "text": "Acknowledgement body."},
    ]

    assert validator.validate_and_heal(structure) == structure


def test_generate_report_applies_custom_formatting_options():
    output = io.BytesIO()
    structure = [{"type": "paragraph", "text": "Body text with custom style."}]

    config = StyleConfig(
        content_font="Arial",
        content_size_pt=10.5,
        line_spacing=2.0,
        margin_inches=0.75
    )

    generate_report(
        structure,
        output,
        style_config=config
    )

    output.seek(0)
    doc = Document(output)
    body_paragraph = next(p for p in doc.paragraphs if p.text == "Body text with custom style.")

    assert body_paragraph.runs[0].font.name == "Arial"
    assert body_paragraph.runs[0].font.size.pt == 10.5
    assert body_paragraph.paragraph_format.line_spacing == 2.0
    assert doc.sections[0].left_margin == Inches(0.75)


def test_visual_style_analysis_accepts_streamlit_like_uploads():
    sample = io.BytesIO()
    sample.name = "reference.docx"

    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run("Styled sample")
    run.font.name = "Arial"
    run.font.size = Pt(13)
    doc.save(sample)
    sample.seek(0)

    profile = StyleAnalyzer().analyze_visual_style(sample)

    assert profile["font_name"] == "Arial"
    assert profile["font_size"] == 13
    assert profile["line_spacing"] == 1.2
