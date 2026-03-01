import os
import pytest
from docx import Document
from src.file_formatting.formatting import generate_report

@pytest.fixture(scope="module")
def sample_docx_path(tmpdir_factory):
    """
    Generates a minimal valid report structure to test the backend `formatting.py` XML generation.
    """
    output_path = str(tmpdir_factory.mktemp("data").join("test_report.docx"))
    
    structure = [
        {"type": "title", "text": "Test Project Title"},
        {"type": "title_page_body", "text": "Submitted by Authors\nUniversity 2025"},
        {"type": "section_header", "text": "ABSTRACT"},
        {"type": "paragraph", "text": "This is the abstract paragraph."},
        {"type": "section_header", "text": "List of Figures"},
        {"type": "lof", "text": "List of Figures"},
        {"type": "section_header", "text": "Contents"},
        {"type": "toc", "text": "Contents"},
        {"type": "chapter", "text": "Introduction"},
        {"type": "paragraph", "text": "Chapter text introducing concepts."},
        {"type": "paragraph", "text": "[Figure 1.1: Architecture Diagram]"},
        {"type": "subheading", "text": "Background"},
        {"type": "paragraph", "text": "This is a background paragraph."}
    ]
    
    generate_report(structure, output_path, style_name="Standard")
    return output_path

def test_docx_can_be_opened(sample_docx_path):
    assert os.path.exists(sample_docx_path)
    doc = Document(sample_docx_path)
    assert doc is not None

def test_sdt_blocks_exist(sample_docx_path):
    """
    Asserts that Native Word SDT elements (w:sdt) were generated. 
    These are used for the dynamic TOC and LOF.
    """
    doc = Document(sample_docx_path)
    
    sdt_count = 0
    # python-docx does not parse SDT blocks as paragraphs natively, 
    # we have to scan the XML element tree directly on the document body.
    for element in doc.element.body.iter():
        if element.tag.endswith("sdt"):
            sdt_count += 1
            
    # We generated 1 TOC and 1 LOF, so there should be exactly two w:sdt blocks
    assert sdt_count == 2, f"Expected 2 SDT blocks for TOC/LOF, found {sdt_count}"

def test_paragraph_styles(sample_docx_path):
    """
    Asserts that the critical Word styles (Heading 1, Heading 2, Caption) exist and were applied.
    """
    doc = Document(sample_docx_path)
    
    styles_used = set(p.style.name for p in doc.paragraphs)
    assert "Heading 1" in styles_used, "Heading 1 style was not applied to chapters."
    assert "Heading 2" in styles_used, "Heading 2 style was not applied to subheadings."
    assert "Caption" in styles_used, "Caption style was not applied to figures."

def test_figure_seq_fields(sample_docx_path):
    """
    Asserts that the [Figure...] placeholder was successfully converted into a 
    w:fldChar block containing the `SEQ Figure \\* ARABIC` instruction.
    """
    doc = Document(sample_docx_path)
    
    seq_found = False
    for p in doc.paragraphs:
        if p.style.name == "Caption":
            # Search paragraph XML for SEQ field instructions
            xml_str = p._element.xml
            if "SEQ Figure" in xml_str and "fldChar" in xml_str:
                seq_found = True
                break
                
    assert seq_found, "Native XML SEQ fields were not injected into the Figure Caption."
