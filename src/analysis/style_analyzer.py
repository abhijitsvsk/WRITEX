import os
import io
from typing import Dict, Any, Counter

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None


class StyleAnalyzer:
    """
    Analyzes sample reports to extract style and formatting guidelines.
    """

    def __init__(self, api_key: str = None):
        self.api_key = api_key
        try:
            from groq import Groq

            self.client = Groq(api_key=api_key) if api_key else None
        except ImportError:
            self.client = None

    def _ocr_image(self, image_data: bytes) -> str:
        """
        Uses Llama 3.2 Vision to extract text from an image.
        """
        if not self.client:
            return ""

        import base64

        try:
            base64_image = base64.b64encode(image_data).decode("utf-8")
            completion = self.client.chat.completions.create(
                model="llama-3.2-11b-vision-preview",
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "Transcribe the text in this image exactly. Return ONLY the text, no conversational filler.",
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{base64_image}"
                                },
                            },
                        ],
                    }
                ],
                temperature=0.1,
                max_tokens=1024,
            )
            return completion.choices[0].message.content or ""
        except Exception as e:
            # Silently fail on OCR errors to avoid breaking the main flow
            print(f"OCR Error: {e}")
            return ""

    def extract_text(self, file_obj, file_path: str) -> str:
        """
        Extracts raw text from a PDF or DOCX file, including OCR for images in PDF.
        """
        text = ""
        try:
            filename = str(file_path).lower()
            if filename.endswith(".pdf"):
                if not PyPDF2:
                    return "Error: PyPDF2 not installed."
                # Handle file_obj differently based on type (bytes vs path)
                reader = PyPDF2.PdfReader(file_obj)

                count = 0
                for page in reader.pages[:15]:  # Analyze first 15 pages for context
                    # 1. Extract Text
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                    # 2. Extract Images (OCR) if text is sparse or if images present
                    # Only do this if we have an API key and haven't processed too many images (cost/latency)
                    if self.client and hasattr(page, "images") and count < 5:
                        for img in page.images:
                            ocr_text = self._ocr_image(img.data)
                            if ocr_text:
                                text += f"\n[OCR Content]: {ocr_text}\n"
                            count += 1

            elif filename.endswith(".docx"):
                if not Document:
                    return "Error: python-docx not installed."
                doc = Document(file_obj)
                for (
                    para
                ) in doc.paragraphs:  # Analyze all paragraphs for structure search
                    text += para.text + "\n"

                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                text += para.text + "\n"

            return text.strip()
        except Exception as e:
            return f"Error extracting text: {e}"

    def extract_specific_sections(self, file_obj, file_path: str) -> Dict[str, str]:
        """
        Extracts specific sections verbatim from the sample report.
        Target sections: Vision, Mission, PEO, PO, PSO, Certificate, Acknowledgment.
        """
        extracted = {}
        text = self.extract_text(file_obj, file_path)

        # Keywords to search for
        targets = {
            "vision": ["vision of the department", "institute vision", "our vision"],
            "mission": [
                "mission of the department",
                "institute mission",
                "our mission",
            ],
            "peo": ["program educational objectives", "peos"],
            "po": ["program outcomes", "pos"],
            "pso": ["program specific outcomes", "psos"],
            "certificate": ["certificate", "bonafide certificate"],
            "acknowledgment": ["acknowledgment", "acknowledgement"],
        }

        # Simple extraction logic: Find heading, take text until next likely heading
        # This is a heuristic.
        lines = text.split("\n")
        current_section = None
        buffer = []

        for line in lines:
            line_clean = line.strip().lower()

            # Check if this line starts a new target section
            found_new = False
            for key, keywords in targets.items():
                for kw in keywords:
                    # Check if line STARTS with keyword (more robust than 'in')
                    if line_clean.startswith(kw):
                        # standard case: short header
                        if len(line_clean) < 50:
                            if current_section and buffer:
                                extracted[current_section] = "\n".join(buffer).strip()
                            current_section = key
                            buffer = []
                            found_new = True
                            break
                        # merged case: "Acknowledgement We wish to..."
                        else:
                            # It's a long line starting with the keyword.
                            # We assume the header is merged with body.
                            if current_section and buffer:
                                extracted[current_section] = "\n".join(buffer).strip()
                            current_section = key
                            buffer = [
                                line[len(kw) :].strip()
                            ]  # Add the rest of the line to buffer
                            found_new = True
                            break
                if found_new:
                    break

            if found_new:
                continue

            # If in a section, accumulate text
            if current_section:
                # Stop if we hit a generic new section
                # stricter checks to avoid false positives in body text
                stops = [
                    "table of contents",
                    "chapter 1",
                    "chapter one",
                    "list of figures",
                    "list of tables",
                    "abbreviations",
                    "declaration",
                ]

                is_stop = False
                if any(x in line_clean for x in stops) and len(line_clean) < 40:
                    is_stop = True

                # Special strict check for single-word headers
                if line_clean in ["abstract", "index", "contents"]:
                    is_stop = True

                if is_stop:
                    extracted[current_section] = "\n".join(buffer).strip()
                    current_section = None
                    buffer = []
                else:
                    buffer.append(line)

        # Capture last section
        if current_section and buffer:
            extracted[current_section] = "\n".join(buffer).strip()

        return extracted

    def analyze_style(self, text: str) -> str:
        """
        Generates a style guide string based on the text.
        (In a real scenario, this could use AI, but for now we use the text itself
        as a few-shot example or just return it as context).
        """
        if not text or "Error" in text:
            return "Use standard academic style."

        # Truncate to avoid overloading prompts (approx 1000 tokens)
        excerpt = text[:3000]

        style_guide = f"""
        STRICT STYLE ADHERENCE REQUIRED.
        
        The user has provided a sample report. You MUST mimic its writing style, tone, and formatting conventions.
        
        Here is an excerpt from the sample report:
        ---
        {excerpt}
        ---
        
        Analyze the excerpt above and adopt the following:
        1. Tone (e.g., passive vs active, formal vs narrative).
        2. Vocabulary level.
        3. Sentence structure complexity.
        
        Write the new content so it feels like it belongs in the SAME document as the excerpt.
        """
        return style_guide

    def detect_structure(self, text: str) -> Dict[str, Any]:
        """
        Analyzes the text to detect structural conventions (e.g., numbered headings).
        """
        structure_config = {
            "numeration": False,  # Default: No numbering (e.g., "Introduction")
            "casing": "Title Case",  # Default
        }

        if not text:
            return structure_config

        # Check for numbered headings (e.g., "1. Introduction", "2. Methodology")
        import re

        # Look for lines starting with "1. ", "2. ", etc., followed by a capital letter
        numbered_matches = re.findall(r"^\d+\.\s+[A-Z]", text, re.MULTILINE)

        if len(numbered_matches) > 2:
            structure_config["numeration"] = True

        return structure_config

    def analyze_visual_style(self, file_path: str) -> Dict[str, Any]:
        """
        Analyzes the visual formatting (font, size, spacing) of a DOCX file.
        Returns the most frequent style attributes.
        """
        style_profile = {
            "font_name": "Times New Roman",  # Default
            "font_size": 12,
            "line_spacing": 1.5,
        }

        if not file_path.endswith(".docx"):
            return style_profile

        if not await_module_availability("docx"):
            return style_profile

        try:
            doc = Document(file_path)

            fonts = []
            sizes = []
            spacings = []

            for para in doc.paragraphs:
                # Line spacing
                if para.paragraph_format.line_spacing:
                    spacings.append(para.paragraph_format.line_spacing)

                for run in para.runs:
                    if run.font.name:
                        fonts.append(run.font.name)
                    if run.font.size:
                        # font.size is in Emu, convert to Pt
                        sizes.append(run.font.size.pt)

            # Determine most frequent
            from collections import Counter

            if fonts:
                style_profile["font_name"] = Counter(fonts).most_common(1)[0][0]

            if sizes:
                # Round to nearest 0.5
                common_size = Counter(sizes).most_common(1)[0][0]
                style_profile["font_size"] = round(common_size * 2) / 2

            if spacings:
                # Average spacing might be better than mode for spacing
                avg_spacing = sum(spacings) / len(spacings)
                style_profile["line_spacing"] = round(avg_spacing, 1)

            return style_profile

        except Exception as e:
            print(f"Error analyzing visual style: {e}")
            return style_profile


def await_module_availability(module_name):
    """
    Checks if a module is available for import.
    """
    import importlib.util

    return importlib.util.find_spec(module_name) is not None
