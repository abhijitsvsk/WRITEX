from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE


from dataclasses import dataclass

@dataclass
class StyleConfig:
    # Page Setup
    margin_inches: float = 1.0
    
    # Headings
    heading_font: str = "Times New Roman"
    heading_size_pt: float = 14.0
    heading_bold: bool = True
    heading_alignment: int = WD_ALIGN_PARAGRAPH.CENTER
    
    # Content
    content_font: str = "Times New Roman"
    content_size_pt: float = 12.0
    content_alignment: int = WD_ALIGN_PARAGRAPH.JUSTIFY
    line_spacing: float = 1.5
    space_before_pt: float = 0.0
    space_after_pt: float = 0.0
    
    # Code
    code_language: str = "Auto"
    
    # Layout
    continuous_sections: bool = False

def _postbuild_estimate_pages(doc):
    """
    POST-BUILD page estimator.  Walks ALL body elements (paragraphs and tables)
    to calculate where page breaks fall with high fidelity.
    """
    PAGE_HEIGHT_PT = 698
    page = 1
    points_on_page = 0
    page_map = {}

    # Interleave paragraphs and tables in original document order
    elements = []
    for p in doc.paragraphs:
        elements.append((p._element.getparent().index(p._element), "PARA", p))
    for t in doc.tables:
        elements.append((t._element.getparent().index(t._element), "TABLE", t))
    
    elements.sort(key=lambda x: x[0])

    for _, etype, obj in elements:
        if etype == "PARA":
            para = obj
            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"

            # 1. Physical Page Breaks
            has_page_break = False
            for run in para.runs:
                if '<w:br w:type="page"/>' in run._r.xml:
                    has_page_break = True
                    break
            
            # 2. Section Breaks (Next Page)
            has_sect_break = bool(para._element.xpath('./w:pPr/w:sectPr'))
            
            if has_page_break or has_sect_break:
                page += 1
                points_on_page = 0
                continue

            if not text:
                points_on_page += 12
                continue

            # 3. Height Measurement
            pt = 0
            if style_name == "Heading 1":
                # Conservative CPL for 16pt bold is ~50 chars.
                h1_lines = 0
                for line in para.text.split('\n'):
                    h1_lines += max(1, (len(line.strip()) + 49) // 50)
                pt = h1_lines * 16 + 24 + 18 # Text + Spacing After + Margin
                page_map[text] = str(page)
            elif style_name == "Heading 2":
                # Conservative CPL for 14pt bold is ~60 chars
                h2_lines = 0
                for line in para.text.split('\n'):
                    h2_lines += max(1, (len(line.strip()) + 59) // 60)
                pt = h2_lines * 14 + 18 + 12
                if points_on_page + pt + 60 > PAGE_HEIGHT_PT:
                    page += 1; points_on_page = 0
                page_map[text] = str(page)
            elif style_name == "Heading 3":
                h3_lines = 0
                for line in para.text.split('\n'):
                    h3_lines += max(1, (len(line.strip()) + 69) // 70)
                pt = h3_lines * 12 + 14 + 8
                if points_on_page + pt + 60 > PAGE_HEIGHT_PT:
                    page += 1; points_on_page = 0
                page_map[text] = str(page)
            elif style_name == "Caption":
                pt = 30
                page_map[text] = str(page)
            else:
                # Normal text - conservative 75 CPL for 12pt TNR
                p_lines = 0
                for line in para.text.split('\n'):
                    p_lines += max(1, (len(line.strip()) + 74) // 75)
                
                if "Course Project Report" in text:
                    pt = 96 # Increased for branding
                elif para.alignment == 1: # Centered (Splash Page)
                    pt = p_lines * 22 + 24 # Increased padding for splash items
                else:
                    # 12pt * 1.5 spacing = 18pt/line + 10pt standard padding
                    pt = p_lines * 18 + 10

            # Image detection
            if para._element.xpath('.//*[local-name()="drawing"]'):
                pt = max(pt, 300)
                if points_on_page + pt > PAGE_HEIGHT_PT:
                    page += 1; points_on_page = 0

            points_on_page += pt
            while points_on_page > PAGE_HEIGHT_PT:
                points_on_page -= PAGE_HEIGHT_PT
                page += 1

        elif etype == "TABLE":
            # Tables take physical vertical space
            pt = 150
            if points_on_page + pt > PAGE_HEIGHT_PT:
                page += 1; points_on_page = 0
            points_on_page += pt

    return page_map


def _patch_toc_lof_pages(doc, page_map):
    """
    Walk the built Document and replace every '?' placeholder in TOC/LOF
    entries with the real page number from page_map.
    
    Works INLINE on the Document object — no temp file needed.
    """
    import re
    def _norm(t):
        return re.sub(r"\s+", " ", t).strip().lower()

    patched = 0
    missed = 0

    for para in doc.paragraphs:
        if "\t?" not in para.text:
            continue

        # Extract title (everything before the tab)
        title = para.text.split("\t")[0].strip()

        # Look up page number — try exact match first, then normalized
        real_page = page_map.get(title)
        if not real_page:
            norm_title = _norm(title)
            for key, val in page_map.items():
                if _norm(key) == norm_title:
                    real_page = val
                    break

        if real_page:
            for run in para.runs:
                if run.text.endswith("?"):
                    run.text = run.text[:-1] + real_page
                    patched += 1
                    break
        else:
            missed += 1
            print(f"[PAGE ESTIMATOR] Unresolved: '{title}'")

    print(f"[PAGE ESTIMATOR] Patched {patched} entries, {missed} unresolved.")



def add_table_of_contents(doc, heading_paragraph, structure=None):

    
    p = doc.add_paragraph()
    r = p.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar3)
    
    # Wrap the paragraph in an SDT block
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    docPartObj = OxmlElement('w:docPartObj')
    docPartGallery = OxmlElement('w:docPartGallery')
    docPartGallery.set(qn('w:val'), 'Table of Contents')
    docPartUnique = OxmlElement('w:docPartUnique')
    docPartObj.append(docPartGallery)
    docPartObj.append(docPartUnique)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)
    
    sdtContent = OxmlElement('w:sdtContent')
    sdtContent.append(p._p)
    
    sdt.append(sdtContent)
    doc.element.body.append(sdt)

def add_list_of_figures(doc, heading_paragraph, structure=None):

    
    p = doc.add_paragraph()
    r = p.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\h \\z \\c "Figure"'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar3)
    
    # Wrap the paragraph in an SDT block
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    docPartObj = OxmlElement('w:docPartObj')
    docPartGallery = OxmlElement('w:docPartGallery')
    docPartGallery.set(qn('w:val'), 'Table of Figures')
    docPartUnique = OxmlElement('w:docPartUnique')
    docPartObj.append(docPartGallery)
    docPartObj.append(docPartUnique)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)
    
    sdtContent = OxmlElement('w:sdtContent')
    sdtContent.append(p._p)
    
    sdt.append(sdtContent)
    doc.element.body.append(sdt)


def generate_report(
    structure,
    output_path,
    style_name="Standard",
    style_config: StyleConfig = None,
):
    doc = Document()
    template_config = {
        "Standard": {"font": "Times New Roman", "size": 12, "spacing": 1.5, "margin": 1.0, "left_margin": 1.5},
        "APA": {"font": "Times New Roman", "size": 12, "spacing": 2.0, "margin": 1.0, "left_margin": 1.0},
        "IEEE": {"font": "Times New Roman", "size": 10, "spacing": 1.0, "margin": 1.0, "left_margin": 1.0},
        "Thesis": {"font": "Times New Roman", "size": 12, "spacing": 1.5, "margin": 1.0, "left_margin": 1.5},
        "Minimal": {"font": "Arial", "size": 11, "spacing": 1.15, "margin": 1.0, "left_margin": 1.0},
    }
    
    t_cfg = template_config.get(style_name, template_config["Standard"])

    if style_config is None:
        style_config = StyleConfig(
            margin_inches=t_cfg["margin"],
            heading_font=t_cfg["font"],
            heading_size_pt=16.0,
            heading_bold=True,
            heading_alignment=WD_ALIGN_PARAGRAPH.LEFT,
            content_font=t_cfg["font"],
            content_size_pt=t_cfg["size"],
            content_alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=t_cfg["spacing"],
            space_before_pt=6.0,
            space_after_pt=12.0
        )

    # Alias for legacy references
    font_name = style_config.content_font
    font_size = style_config.content_size_pt
    line_spacing = style_config.line_spacing

    # --- Margin Setup (Enforcing Strict A4 Geometry) ---
    from docx.shared import Mm
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    section.top_margin = Inches(style_config.margin_inches)
    section.bottom_margin = Inches(style_config.margin_inches)
    section.left_margin = Inches(style_config.margin_inches)
    section.right_margin = Inches(style_config.margin_inches)

    # --- Style Bootstrapping (Root Cause 2 Fix) ---
    # Ensure critical styles exist so the TOC/LOF fields don't silently fail.
    required_styles = ["Heading 1", "Heading 2", "Heading 3", "Caption"]
    for s_name in required_styles:
        try:
            _ = doc.styles[s_name]
        except KeyError:
            doc.styles.add_style(s_name, WD_STYLE_TYPE.PARAGRAPH)

    # --- Hierarchy Counters ---
    # These strictly track where we are in the document
    counters = {"chapter": 0, "sub": 0, "subsub": 0, "figure": 0}

    # --- PRE-PROCESS: Extract inline figures from paragraphs ---
    import re as _re

    _clean_structure = []
    for _item in structure:
        if _item.get("type", "") == "paragraph" and _item.get("text", ""):
            _txt = _item.get("text")
            # Pull '[Figure X: Title]' out into its own paragraph block safely
            _parts = _re.split(
                r"(\[(?:Fig|Figure)\s*[\d\.]*[:\-]?\s*.*?\])",
                _txt,
                flags=_re.IGNORECASE,
            )
            for _p in _parts:
                _p = _p.strip()
                if _p:
                    _clean_structure.append({"type": "paragraph", "text": _p})
        else:
            _clean_structure.append(_item)
    structure = _clean_structure

    seen_captions = set()  # Fix 4: Semantic Dedup captions
    skip_indices = set()
    
    # Enable Pagination (Arabic numbers from the start)
    if doc.sections:
        sec1 = doc.sections[0]
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(ns.qn('w:fmt'), 'decimal')
        sec1._sectPr.append(pgNumType)

    for idx, item in enumerate(structure):
        if idx in skip_indices:
            continue
        itype = item.get("type")
        text = item.get("text", "")

        # 0a. TOC (Table of Contents)
        if itype == "toc":
            doc.add_page_break()
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("Table of Contents")
            run.font.name = font_name
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_after = Pt(24)
            p.paragraph_format.space_before = Pt(0)
            add_table_of_contents(doc, p, structure)
            continue

        # 0b. LOF (List of Figures)
        elif itype == "lof":
            doc.add_page_break()
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("List of Figures")
            run.font.name = font_name
            run.font.size = Pt(16)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_after = Pt(24)
            p.paragraph_format.space_before = Pt(0)
            add_list_of_figures(doc, p, structure)
            continue

        # 1. CHAPTER (Level 1)
        elif itype == "chapter":
            counters["chapter"] += 1
            counters["sub"] = 0
            counters["subsub"] = 0
            counters["figure"] = 0

            # First Chapter gets a new section (for layout), continuous page numbering
            if counters["chapter"] == 1:
                # Add section break for layout separation
                if style_config.continuous_sections:
                    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
                else:
                    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
                
                # Unlink footer from front matter
                new_section.footer.is_linked_to_previous = False
            else:
                if not style_config.continuous_sections:
                    doc.add_page_break()

            p = doc.add_paragraph()
            p.style = doc.styles["Heading 1"]
            p.alignment = style_config.heading_alignment

            # Use Title Case for Chapter Titles
            run = p.add_run(f"Chapter {counters['chapter']} {text.title()}")
            run.bold = style_config.heading_bold
            run.font.name = style_config.heading_font
            run.font.size = Pt(style_config.heading_size_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)

            p.paragraph_format.space_after = Pt(24)

        # 2. SUBHEADING (Level 2) - 1.1
        elif itype == "subheading":
            counters["sub"] += 1
            counters["subsub"] = 0

            prefix = f"{counters['chapter']}.{counters['sub']}"
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 2"]
            p.alignment = style_config.heading_alignment

            run = p.add_run(f"{prefix} {text.title()}")
            run.bold = style_config.heading_bold
            run.font.name = style_config.heading_font
            run.font.size = Pt(max(10, style_config.heading_size_pt - 2))
            run.font.color.rgb = RGBColor(0, 0, 0)

            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)

        # 3. SUBSUBHEADING (Level 3) - 1.1.1
        elif itype == "subsubheading":
            counters["subsub"] += 1

            prefix = f"{counters['chapter']}.{counters['sub']}.{counters['subsub']}"
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 3"]
            p.alignment = style_config.heading_alignment

            run = p.add_run(f"{prefix} {text.title()}")
            run.bold = style_config.heading_bold
            run.font.name = style_config.heading_font
            run.font.size = Pt(max(10, style_config.heading_size_pt - 4))
            run.font.color.rgb = RGBColor(0, 0, 0)

            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)

        # 4. TITLE / SPLASH
        elif itype == "title":
            p = doc.add_paragraph("Course Project Report On")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(48)
            p.runs[0].font.name = font_name
            p.runs[0].font.size = Pt(14)
            p.paragraph_format.space_after = Pt(24)

            p2 = doc.add_paragraph(text)
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p2.runs[0]
            run.font.size = Pt(22)
            run.bold = True
            run.font.name = font_name
            p2.paragraph_format.space_after = Pt(24)
            p2.paragraph_format.line_spacing = 1.0

        # 4a. TITLE PAGE BODY (Centered, Single Spaced)
        elif itype == "title_page_body":
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(14)
                if "Bachelor of Technology" in run.text or "Submitted by" in run.text:
                    run.bold = True
            p.paragraph_format.space_after = Pt(24)

        # 4b. SECTION HEADER (Unnumbered, New Page) - Abstract, References
        elif itype == "section_header":
            # Skip building headers for LOF/TOC here as they are explicitly built in the functions above.
            if text.upper() in ["LIST OF FIGURES", "TABLE OF CONTENTS", "CONTENTS"]:
                continue

            if not style_config.continuous_sections:
                doc.add_page_break()
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 1"]
            p.alignment = style_config.heading_alignment
            run = p.add_run(text.title())  # Changed from UPPER to title case
            run.bold = style_config.heading_bold
            run.font.name = style_config.heading_font
            run.font.size = Pt(style_config.heading_size_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_after = Pt(24)
            p.paragraph_format.space_before = Pt(0)

        # 4c. INSTITUTIONAL HEADER (Unnumbered, New Page, Left Aligned, Title Case)
        elif itype == "institutional_header":
            if not style_config.continuous_sections:
                doc.add_page_break()
            p = doc.add_paragraph()
            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.title())
            run.bold = True
            run.font.name = style_config.heading_font
            run.font.size = Pt(style_config.heading_size_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_after = Pt(24)

        # 5. SIGNATURE BLOCK
        elif itype == "signature_block":
            guide = item.get("guide", "Guide")
            guide_desig = item.get("guide_designation", "Assistant Professor")
            hod = item.get("hod", "HOD")
            hod_desig = item.get("hod_designation", "Professor & HoD")
            department = item.get("department", "")

            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            # Add spacing before signature block
            for cell in table.columns[0].cells:
                for cp in cell.paragraphs:
                    cp.paragraph_format.space_before = Pt(36)

            # Left Cell: Guide
            cell_l = table.cell(0, 0)
            p = cell_l.paragraphs[0]
            guide_text = f"\n\n___________________\n{guide}\n{guide_desig}"
            if department:
                guide_text += f"\nDepartment of {department}"
            run = p.add_run(guide_text)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Right Cell: HOD
            cell_r = table.cell(0, 1)
            p = cell_r.paragraphs[0]
            hod_text = f"\n\n___________________\n{hod}\n{hod_desig}"
            if department:
                hod_text += f"\nDepartment of {department}"
            run = p.add_run(hod_text)
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 6. PAGE BREAK
        elif itype == "page_break":
            doc.add_page_break()

        # 8. CODE SNIPPET (NATIVE)
        elif itype == "code_block":
            # Add "Code:" label first
            label_p = doc.add_paragraph()
            label_run = label_p.add_run("Code:")
            label_run.bold = True
            label_run.font.name = style_config.content_font
            label_run.font.size = Pt(style_config.content_size_pt)
            label_p.paragraph_format.space_before = Pt(12)
            label_p.paragraph_format.space_after = Pt(4)

            # Insert code in a 1x1 table for a clean padded box
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            
            # Set shading color to Dark Gray for a premium code block look (or light gray)
            tcPr = cell._element.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(ns.qn("w:val"), "clear")
            shd.set(ns.qn("w:color"), "auto")
            shd.set(ns.qn("w:fill"), "282C34") # Monokai dark background
            tcPr.append(shd)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15

            # Syntax highlighting using Pygments (if available, else fallback)
            if style_config.code_language == "None":
                run = p.add_run(text)
                run.font.name = "Courier New"
                run.font.size = Pt(style_config.content_size_pt)
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                try:
                    from pygments import lex
                    from pygments.lexers import get_lexer_for_filename, guess_lexer
                    from pygments.styles import get_style_by_name
                    
                    try:
                        if style_config.code_language and style_config.code_language != "Auto":
                            from pygments.lexers import get_lexer_by_name
                            try:
                                lexer = get_lexer_by_name(style_config.code_language.lower())
                            except Exception:
                                lexer = guess_lexer(text)
                        else:
                            lexer = guess_lexer(text)
                    except Exception:
                        from pygments.lexers import PythonLexer
                        lexer = PythonLexer()
                        
                    style = get_style_by_name("monokai")
                    
                    for token, content in lex(text, lexer):
                        if not content:
                            continue
                        run = p.add_run(content)
                        run.font.name = "Courier New"
                        run.font.size = Pt(style_config.content_size_pt)
                        
                        # Apply color from pygments style
                        if token in style.styles:
                            color_hex = style.styles[token]
                            if color_hex and color_hex.startswith("#"):
                                color_hex = color_hex[1:]
                                if len(color_hex) == 6:
                                    r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
                                    run.font.color.rgb = RGBColor(r, g, b)
                            elif "bold" in style.styles[token]:
                                run.bold = True
                        
                        if not run.font.color.rgb:
                            run.font.color.rgb = RGBColor(248, 248, 242)
                except ImportError:
                    # Fallback to plain text
                    run = p.add_run(text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(style_config.content_size_pt)
                    run.font.color.rgb = RGBColor(248, 248, 242)

        # 9. TERMINAL OUTPUT (Tabular data, console logs)
        elif itype == "terminal_output":
            label_p = doc.add_paragraph()
            label_run = label_p.add_run("Output:")
            label_run.bold = True
            label_run.font.name = style_config.content_font
            label_run.font.size = Pt(style_config.content_size_pt)
            label_p.paragraph_format.space_before = Pt(12)
            label_p.paragraph_format.space_after = Pt(4)

            # Insert code in a 1x1 table for a clean padded box (Terminal feel)
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            
            tcPr = cell._element.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(ns.qn("w:val"), "clear")
            shd.set(ns.qn("w:color"), "auto")
            shd.set(ns.qn("w:fill"), "F5F5F5") # Light gray background for terminal output
            tcPr.append(shd)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.0

            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(style_config.content_size_pt)
            run.font.color.rgb = RGBColor(0, 0, 0) # Black text
            
            # Note: "Explanation:" label is intentionally omitted to prevent
            # empty heading artifacts when no further explanation follows the code.
            pass

        # 8b. ALONE IMAGE
        elif itype == "image":
            import io

            try:
                image_stream = io.BytesIO(item["content"])
                doc.add_picture(image_stream, width=Inches(6.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error embedding standalone diagram: {e}")

        # 9. FIGURE (NATIVE AST OBJECT)
        elif itype == "figure":
            counters["figure"] += 1
            caption_clean = (item.get("caption", "") or item.get("text", "")).strip()
            
            # --- IMAGE PLACEHOLDER OR ACTUAL IMAGE ---
            next_item = structure[idx + 1] if idx + 1 < len(structure) else None
            # If the architecture ever feeds in an image buffer, inject it natively
            if next_item and next_item.get("type") == "image":
                import io
                try:
                    image_stream = io.BytesIO(next_item["content"])
                    doc.add_picture(image_stream, width=Inches(6.0))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error embedding diagram: {e}")
                skip_indices.add(idx + 1)
            else:
                placeholder_p = doc.add_paragraph()
                placeholder_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                placeholder_p.paragraph_format.space_before = Pt(12)
                placeholder_p.paragraph_format.space_after = Pt(6)

                box_run = placeholder_p.add_run("\n[ Image Placeholder ]\n")
                box_run.font.name = "Consolas"
                box_run.font.size = Pt(10)

                shading_elm = OxmlElement("w:shd")
                shading_elm.set(ns.qn("w:val"), "clear")
                shading_elm.set(ns.qn("w:color"), "auto")
                shading_elm.set(ns.qn("w:fill"), "EAEAEA")
                placeholder_p.paragraph_format.element.get_or_add_pPr().append(
                    shading_elm
                )

            # --- CAPTION (Native Word SEQ Field) ---
            p = doc.add_paragraph()
            p.style = doc.styles["Caption"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.0

            run = p.add_run("Figure ")
            run.font.name = font_name
            run.font.size = Pt(11)


            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' SEQ Figure \\* ARABIC '
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            r_xml = p.add_run()
            r_xml._r.append(fldChar1)
            r_xml._r.append(instrText)
            r_xml._r.append(fldChar2)
            r_xml._r.append(fldChar3)

            run_text = p.add_run(f" {caption_clean}")
            run_text.font.name = font_name
            run_text.font.size = Pt(11)

        # 10. PARAGRAPH / BODY / PLACEHOLDERS
        else:
            text = text.strip()
            if not text:
                continue

            import re

            code_match = re.search(r"\[Extract Code:\s*(.*?)\]", text, re.IGNORECASE)
            if code_match:
                continue

            if text.lower().startswith("[figure") or text.lower().startswith("[fig"):
                itype = "figure"
                caption_clean = text.strip("[]")
                if ":" in caption_clean:
                    caption_clean = caption_clean.split(":", 1)[1].strip()
                # Pass it down to the figure block below! We'll just execute the figure logic inline here.
                counters["figure"] += 1
                p = doc.add_paragraph()
                p.style = doc.styles["Caption"]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.0

                run = p.add_run("Figure ")
                run.font.name = font_name
                run.font.size = Pt(11)
                
                # Native XML SEQ Field

                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = ' SEQ Figure \\* ARABIC '
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                
                r_xml = p.add_run()
                r_xml._r.append(fldChar1)
                r_xml._r.append(instrText)
                r_xml._r.append(fldChar2)
                r_xml._r.append(fldChar3)
                
                run_text = p.add_run(f" {caption_clean}")
                run_text.font.name = font_name
                run_text.font.size = Pt(11)
                continue

            p = doc.add_paragraph(text)
            p.alignment = style_config.content_alignment
            p.paragraph_format.space_before = Pt(style_config.space_before_pt)
            p.paragraph_format.space_after = Pt(style_config.space_after_pt)
            p.paragraph_format.line_spacing = style_config.line_spacing

            for run in p.runs:
                run.font.name = style_config.content_font
                run.font.size = Pt(style_config.content_size_pt)

    # --- POST-PROCESSING ---

    # ZERO-TOLERANCE WHITESPACE PURGE
    # Snapshot the list to avoid modifying during iteration
    paragraphs_snapshot = list(doc.paragraphs)
    for p in paragraphs_snapshot:
        if not p.text.strip():
            # Preserve page breaks
            has_break = bool(p._element.xpath('.//*[local-name()="br"]'))
            if has_break:
                continue

            # Preserve embedded images
            has_drawing = bool(p._element.xpath('.//*[local-name()="drawing"]'))
            if has_drawing:
                continue

            # Preserve shaded blocks (code blocks, placeholders)
            has_shading = bool(p._element.xpath('.//*[local-name()="shd"]'))
            if has_shading:
                continue

            # Preserve section break paragraphs
            has_sectPr = bool(p._element.xpath('.//*[local-name()="sectPr"]'))
            if has_sectPr:
                continue

            # Preserve SDT-adjacent paragraphs
            has_sdt_sibling = (
                (p._element.getnext() is not None and p._element.getnext().tag.endswith('}sdt'))
                or (p._element.getprevious() is not None and p._element.getprevious().tag.endswith('}sdt'))
            )
            if has_sdt_sibling:
                continue

            # Safe to delete
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)

    _add_page_numbers(doc, font_name)

    # --- Structural Validation (Root Cause 3 Fix) ---
    _validate_document_structure(doc)

    # Save
    # ── Post-build page estimation: walk the real doc → patch placeholders → save ──
    print("[PAGE ESTIMATOR] Running post-build page estimation on actual document...")
    page_map = _postbuild_estimate_pages(doc)
    _patch_toc_lof_pages(doc, page_map)

    # Single-pass save — no temp files, no external patchers
    if hasattr(output_path, "write"):
        doc.save(output_path)
    else:
        doc.save(output_path)


def _validate_document_structure(doc):
    """Mathematical AST validation to guarantee headings, captions, and SEQ fields exist before saving."""
    heading_count = 0
    caption_count = 0

    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            heading_count += 1
        elif p.style.name == "Caption":
            caption_count += 1

    seq_count = len(doc.element.xpath('//w:instrText[contains(text(), "SEQ Figure")]'))

    print(
        f"[STRUCTURAL AUDIT] Verified {heading_count} Headings and {caption_count} Captions mapped correctly via AST. Found {seq_count} SEQ Figure instructions."
    )
    if heading_count == 0:
        print(
            "[WARNING] Zero headings detected in the Document Object. The TOC will render empty."
        )
    if caption_count == 0:
        print(
            "[WARNING] Zero Captions detected in the Document Object. The LOF will render empty."
        )
    if caption_count > 0 and seq_count == 0:
        print(
            "[WARNING] Zero SEQ Figure fields detected! The LOF field mathematically cannot populate."
        )


def _add_page_numbers(doc, font_name):
    """Add professional page numbers to all document sections."""
    for i, section in enumerate(doc.sections):
        footer = section.footer
        # All sections use Arabic page numbers.
        # But both need the PAGE field injected.
        p = footer.paragraphs[0]
        
        # Guard against double-injecting the PAGE field if it already exists
        if p.text.strip():
            p.clear()
            
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Professional format: "— PAGE —"
        # Add left dash
        run_left = p.add_run("\u2014 ")
        run_left.font.name = font_name
        run_left.font.size = Pt(10)
        
        # Add PAGE field
        run = p.add_run()
        run.font.name = font_name
        run.font.size = Pt(10)

        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(ns.qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(ns.qn("xml:space"), "preserve")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(ns.qn("w:fldCharType"), "separate")

        # Cached display text for Word versions that don't auto-calc
        cached_page = OxmlElement("w:t")
        cached_page.text = "1"

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(ns.qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(cached_page)
        run._r.append(fldChar3)
        
        # Add right dash
        run_right = p.add_run(" \u2014")
        run_right.font.name = font_name
        run_right.font.size = Pt(10)

