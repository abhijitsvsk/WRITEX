import re

with open(r"d:\Z_shared\writex\src\file_formatting\formatting.py", "r", encoding="utf-8") as f:
    content = f.read()

# 1. Update add_table_of_contents
toc_replace = """def add_table_of_contents(doc, heading_paragraph, structure=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    p = doc.add_paragraph()
    r = p.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\\\o "1-3" \\\\h \\\\z \\\\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar3)
    
    # Wrap the paragraph in an SDT block
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    docPartObj = OxmlElement('w:docPartObj')
    docPartGallery = OxmlElement('w:docPartGallery')
    docPartGallery.set(qn('w:val'), 'Table of Contents')
    docPartUnique = OxmlElement('w:docPartUnique')
    docPartObj.append(docPartGallery)
    docPartObj.append(docPartUnique)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)
    
    sdtContent = OxmlElement('w:sdtContent')
    sdtContent.append(p._p)
    
    sdt.append(sdtContent)
    doc.element.body.append(sdt)
"""
content = re.sub(
    r'def add_table_of_contents\(doc, heading_paragraph, structure=None\):.*?def add_list_of_figures',
    lambda m: toc_replace + '\ndef add_list_of_figures',
    content,
    flags=re.DOTALL
)

# 2. Update add_list_of_figures
lof_replace = """def add_list_of_figures(doc, heading_paragraph, structure=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    p = doc.add_paragraph()
    r = p.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\\\h \\\\z \\\\c "Figure"'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    r._r.append(fldChar3)
    
    # Wrap the paragraph in an SDT block
    sdt = OxmlElement('w:sdt')
    sdtPr = OxmlElement('w:sdtPr')
    docPartObj = OxmlElement('w:docPartObj')
    docPartGallery = OxmlElement('w:docPartGallery')
    docPartGallery.set(qn('w:val'), 'Table of Figures')
    docPartUnique = OxmlElement('w:docPartUnique')
    docPartObj.append(docPartGallery)
    docPartObj.append(docPartUnique)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)
    
    sdtContent = OxmlElement('w:sdtContent')
    sdtContent.append(p._p)
    
    sdt.append(sdtContent)
    doc.element.body.append(sdt)
"""
content = re.sub(
    r'def add_list_of_figures\(doc, heading_paragraph, structure=None\):.*?def generate_report',
    lambda m: lof_replace + '\n\ndef generate_report',
    content,
    flags=re.DOTALL
)

# 3. Update paragraph logic to catch [Figure...] strings and format them as Captions
figure_catch = """        # 10. PARAGRAPH / BODY / PLACEHOLDERS
        else:
            text = text.strip()
            if not text:
                continue

            import re

            code_match = re.search(r"\\[Extract Code:\\s*(.*?)\\]", text, re.IGNORECASE)
            if code_match:
                continue

            if text.lower().startswith("[figure") or text.lower().startswith("[fig"):
                itype = "figure"
                caption_clean = text.strip("[]")
                if ":" in caption_clean:
                    caption_clean = caption_clean.split(":", 1)[1].strip()
                # Pass it down to the figure block below! We'll just execute the figure logic inline here.
                counters["figure"] += 1
                p = doc.add_paragraph()
                p.style = doc.styles["Caption"]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(12)
                p.paragraph_format.line_spacing = 1.0

                run = p.add_run("Figure ")
                run.font.name = font_name
                run.font.size = Pt(11)
                
                # Native XML SEQ Field
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = ' SEQ Figure \\\\* ARABIC '
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                
                r_xml = p.add_run()
                r_xml._r.append(fldChar1)
                r_xml._r.append(instrText)
                r_xml._r.append(fldChar2)
                r_xml._r.append(fldChar3)
                
                run_text = p.add_run(f" {caption_clean}")
                run_text.font.name = font_name
                run_text.font.size = Pt(11)
                continue

            p = doc.add_paragraph(text)"""

content = re.sub(
    r'        # 10\. PARAGRAPH / BODY / PLACEHOLDERS\s*else:\s*text = text\.strip\(\)\s*if not text:\s*continue.*?(?=p = doc\.add_paragraph\(text\))',
    lambda m: figure_catch + '\n            ',
    content,
    flags=re.DOTALL
)

# 4. Update the actual "figure" logic for AST objects
figure_ast_catch = """            # --- CAPTION (Native Word SEQ Field) ---
            p = doc.add_paragraph()
            p.style = doc.styles["Caption"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.0

            run = p.add_run("Figure ")
            run.font.name = font_name
            run.font.size = Pt(11)

            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = ' SEQ Figure \\\\* ARABIC '
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')

            r_xml = p.add_run()
            r_xml._r.append(fldChar1)
            r_xml._r.append(instrText)
            r_xml._r.append(fldChar2)
            r_xml._r.append(fldChar3)

            run_text = p.add_run(f" {caption_clean}")
            run_text.font.name = font_name
            run_text.font.size = Pt(11)"""
            
content = re.sub(
    r'            # --- CAPTION \(Fully Deterministic — No Word Field Dependencies\) ---.*?run\.font\.size = Pt\(11\)',
    lambda m: figure_ast_catch,
    content,
    flags=re.DOTALL
)

with open(r"d:\Z_shared\writex\src\file_formatting\formatting.py", "w", encoding="utf-8") as f:
    f.write(content)

print("formatting.py patched!")
