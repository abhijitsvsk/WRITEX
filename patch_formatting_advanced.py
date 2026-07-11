import re

with open(r"d:\Z_shared\writex\src\file_formatting\formatting.py", "r", encoding="utf-8") as f:
    content = f.read()

# Inject StyleConfig dataclass at the top
style_config_code = """from dataclasses import dataclass

@dataclass
class StyleConfig:
    # Page Setup
    margin_inches: float = 1.0
    
    # Headings
    heading_font: str = "Times New Roman"
    heading_size_pt: float = 14.0
    heading_bold: bool = True
    heading_alignment: int = WD_ALIGN_PARAGRAPH.CENTER
    
    # Content
    content_font: str = "Times New Roman"
    content_size_pt: float = 12.0
    content_alignment: int = WD_ALIGN_PARAGRAPH.JUSTIFY
    line_spacing: float = 1.5
    space_before_pt: float = 0.0
    space_after_pt: float = 0.0

def _postbuild_estimate_pages(doc):"""

content = content.replace("def _postbuild_estimate_pages(doc):", style_config_code)

# Refactor generate_report signature
gen_sig_old = """def generate_report(
    structure,
    output_path,
    style_name="Standard",
    custom_font=None,
    custom_size=None,
    custom_spacing=None,
    custom_margin=None,
):"""

gen_sig_new = """def generate_report(
    structure,
    output_path,
    style_name="Standard",
    style_config: StyleConfig = None,
):"""

content = content.replace(gen_sig_old, gen_sig_new)

# Update config loading
cfg_load_old = """    t_cfg = template_config.get(style_name, template_config["Standard"])

    # --- Standard Style Config ---
    font_name = custom_font or t_cfg["font"]
    font_size = custom_size or t_cfg["size"]
    line_spacing = custom_spacing or t_cfg["spacing"]

    # --- Margin Setup (Enforcing Strict A4 Geometry) ---
    from docx.shared import Mm
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    if custom_margin:
        section.top_margin = Inches(custom_margin)
        section.bottom_margin = Inches(custom_margin)
        section.left_margin = Inches(custom_margin)
        section.right_margin = Inches(custom_margin)
    else:
        section.top_margin = Inches(t_cfg["margin"])
        section.bottom_margin = Inches(t_cfg["margin"])
        section.left_margin = Inches(t_cfg["left_margin"])
        section.right_margin = Inches(t_cfg["margin"])"""

cfg_load_new = """    t_cfg = template_config.get(style_name, template_config["Standard"])

    if style_config is None:
        style_config = StyleConfig(
            margin_inches=t_cfg["margin"],
            heading_font=t_cfg["font"],
            heading_size_pt=16.0,
            heading_bold=True,
            heading_alignment=WD_ALIGN_PARAGRAPH.LEFT,
            content_font=t_cfg["font"],
            content_size_pt=t_cfg["size"],
            content_alignment=WD_ALIGN_PARAGRAPH.LEFT,
            line_spacing=t_cfg["spacing"],
            space_before_pt=6.0,
            space_after_pt=12.0
        )

    # Alias for legacy references
    font_name = style_config.content_font
    font_size = style_config.content_size_pt
    line_spacing = style_config.line_spacing

    # --- Margin Setup (Enforcing Strict A4 Geometry) ---
    from docx.shared import Mm
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    
    section.top_margin = Inches(style_config.margin_inches)
    section.bottom_margin = Inches(style_config.margin_inches)
    section.left_margin = Inches(style_config.margin_inches)
    section.right_margin = Inches(style_config.margin_inches)"""

content = content.replace(cfg_load_old, cfg_load_new)

# Update CHAPTER styling
chapter_old = """            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Use Title Case for Chapter Titles
            run = p.add_run(f"Chapter {counters['chapter']} {text.title()}")
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)"""

chapter_new = """            p.style = doc.styles["Heading 1"]
            p.alignment = style_config.heading_alignment

            # Use Title Case for Chapter Titles
            run = p.add_run(f"Chapter {counters['chapter']} {text.title()}")
            run.bold = style_config.heading_bold
            run.font.name = style_config.heading_font
            run.font.size = Pt(style_config.heading_size_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)"""
content = content.replace(chapter_old, chapter_new)

# Update SUBHEADING styling
sub_old = """            p.style = doc.styles["Heading 2"]

            run = p.add_run(f"{prefix} {text.title()}")
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)"""
sub_new = """            p.style = doc.styles["Heading 2"]
            p.alignment = style_config.heading_alignment

            run = p.add_run(f"{prefix} {text.title()}")
            run.bold = style_config.heading_bold
            run.font.name = style_config.heading_font
            run.font.size = Pt(max(10, style_config.heading_size_pt - 2))
            run.font.color.rgb = RGBColor(0, 0, 0)"""
content = content.replace(sub_old, sub_new)

# Update SUBSUBHEADING styling
subsub_old = """            p.style = doc.styles["Heading 3"]

            run = p.add_run(f"{prefix} {text.title()}")
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)"""
subsub_new = """            p.style = doc.styles["Heading 3"]
            p.alignment = style_config.heading_alignment

            run = p.add_run(f"{prefix} {text.title()}")
            run.bold = style_config.heading_bold
            run.font.name = style_config.heading_font
            run.font.size = Pt(max(10, style_config.heading_size_pt - 4))
            run.font.color.rgb = RGBColor(0, 0, 0)"""
content = content.replace(subsub_old, subsub_new)

# Update SECTION HEADER styling
sect_old = """            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.title())  # Changed from UPPER to title case
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0, 0, 0)"""
sect_new = """            p.style = doc.styles["Heading 1"]
            p.alignment = style_config.heading_alignment
            run = p.add_run(text.title())  # Changed from UPPER to title case
            run.bold = style_config.heading_bold
            run.font.name = style_config.heading_font
            run.font.size = Pt(style_config.heading_size_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)"""
content = content.replace(sect_old, sect_new)

# Update INSTITUTIONAL HEADER
inst_old = """            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.title())
            run.bold = True
            run.font.name = font_name
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0, 0, 0)"""
inst_new = """            p.style = doc.styles["Heading 1"]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text.title())
            run.bold = True
            run.font.name = style_config.heading_font
            run.font.size = Pt(style_config.heading_size_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)"""
content = content.replace(inst_old, inst_new)

# Update CODE BLOCK formatting
code_old = """        # 8. CODE SNIPPET (NATIVE)
        elif itype == "code_block":
            # Add "Code:" label first
            label_p = doc.add_paragraph()
            label_run = label_p.add_run("Code:")
            label_run.bold = True
            label_run.font.name = font_name
            label_run.font.size = Pt(font_size)
            label_p.paragraph_format.space_before = Pt(12)
            label_p.paragraph_format.space_after = Pt(4)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Simple gray "code block" background via shading xml element
            shading_elm = OxmlElement("w:shd")
            shading_elm.set(ns.qn("w:val"), "clear")
            shading_elm.set(ns.qn("w:color"), "auto")
            shading_elm.set(ns.qn("w:fill"), "F0F0F0")  # Light gray
            p.paragraph_format.element.get_or_add_pPr().append(shading_elm)

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.0  # Code is single-spaced
            p.paragraph_format.left_indent = Inches(0.25)

            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0, 0, 0)"""
            
code_new = """        # 8. CODE SNIPPET (NATIVE)
        elif itype == "code_block":
            # Add "Code:" label first
            label_p = doc.add_paragraph()
            label_run = label_p.add_run("Code:")
            label_run.bold = True
            label_run.font.name = style_config.content_font
            label_run.font.size = Pt(style_config.content_size_pt)
            label_p.paragraph_format.space_before = Pt(12)
            label_p.paragraph_format.space_after = Pt(4)

            # Insert code in a 1x1 table for a clean padded box
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            
            # Set shading color to Dark Gray for a premium code block look (or light gray)
            tcPr = cell._element.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(ns.qn("w:val"), "clear")
            shd.set(ns.qn("w:color"), "auto")
            shd.set(ns.qn("w:fill"), "282C34") # Monokai dark background
            tcPr.append(shd)

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15

            # Syntax highlighting using Pygments (if available, else fallback)
            try:
                from pygments import lex
                from pygments.lexers import get_lexer_for_filename, guess_lexer
                from pygments.styles import get_style_by_name
                
                try:
                    lexer = guess_lexer(text)
                except Exception:
                    from pygments.lexers import PythonLexer
                    lexer = PythonLexer()
                    
                style = get_style_by_name("monokai")
                
                for token, content in lex(text, lexer):
                    if not content:
                        continue
                    run = p.add_run(content)
                    run.font.name = "Courier New"
                    run.font.size = Pt(style_config.content_size_pt) # Adhere to constraint!
                    
                    # Apply color from pygments style
                    if token in style.styles:
                        color_hex = style.styles[token]
                        if color_hex and color_hex.startswith("#"):
                            color_hex = color_hex[1:]
                            if len(color_hex) == 6:
                                r, g, b = tuple(int(color_hex[i:i+2], 16) for i in (0, 2, 4))
                                run.font.color.rgb = RGBColor(r, g, b)
                        elif "bold" in style.styles[token]:
                            run.bold = True
                    
                    if not run.font.color.rgb:
                        run.font.color.rgb = RGBColor(248, 248, 242) # Default text color (off-white)
            except ImportError:
                # Fallback to plain text
                run = p.add_run(text)
                run.font.name = "Courier New"
                run.font.size = Pt(style_config.content_size_pt)
                run.font.color.rgb = RGBColor(248, 248, 242)"""
content = content.replace(code_old, code_new)

# Update REGULAR PARAGRAPH styling
para_old = """            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = line_spacing

            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)"""
para_new = """            p.alignment = style_config.content_alignment
            p.paragraph_format.space_before = Pt(style_config.space_before_pt)
            p.paragraph_format.space_after = Pt(style_config.space_after_pt)
            p.paragraph_format.line_spacing = style_config.line_spacing

            for run in p.runs:
                run.font.name = style_config.content_font
                run.font.size = Pt(style_config.content_size_pt)"""
content = content.replace(para_old, para_new)

with open(r"d:\Z_shared\writex\src\file_formatting\formatting.py", "w", encoding="utf-8") as f:
    f.write(content)
print("Updated formatting.py")
