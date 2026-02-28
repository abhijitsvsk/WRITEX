from docx import Document

def create_seed():
    doc = Document()
    doc.add_heading('Project Blast', 0)
    
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('This is a deterministic self-healing automation system named Blast.')
    
    doc.add_heading('Architecture', level=1)
    doc.add_paragraph('Blast is composed of Blueprint, Architect, and Trigger directives.')
    
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('The system is robust and efficient.')
    
    doc.save('d:/writex/seed.docx')
    print("Created d:/writex/seed.docx")

if __name__ == "__main__":
    create_seed()
